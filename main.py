from __future__ import annotations

import csv
import hmac
import io
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Cm, RGBColor
from copy import deepcopy
from datetime import date, datetime
from pathlib import Path
from typing import Any, Dict, List

import pandas as pd
import streamlit as st
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


# -----------------------------
# CONFIG APP
# -----------------------------
st.set_page_config(
    page_title="Audit Certificazione ISO",
    page_icon="📋",
    layout="wide",
)

NAVY = "#16324F"
NAVY_2 = "#0E2236"
AZURE = "#4EA1FF"
LIGHT = "#F6F9FC"
SOFT = "#EAF2FB"
TEXT = "#122230"
BORDER = "rgba(18,34,48,0.12)"

st.markdown(
    f"""
<style>
:root {{
  --navy: {NAVY};
  --navy2: {NAVY_2};
  --azure: {AZURE};
  --light: {LIGHT};
  --soft: {SOFT};
  --text: {TEXT};
  --border: {BORDER};
}}

.stApp {{
  background: linear-gradient(180deg, var(--light) 0%, #ffffff 60%);
  color: var(--text);
}}

/* SIDEBAR */
section[data-testid="stSidebar"] {{
  background: linear-gradient(180deg, var(--navy) 0%, var(--navy2) 100%);
  border-right: 1px solid rgba(255,255,255,0.08);
}}

section[data-testid="stSidebar"] .stMarkdown,
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] .stCaption,
section[data-testid="stSidebar"] .stSubheader,
section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] span {{
  color: #F8FBFF !important;
}}

section[data-testid="stSidebar"] .stSubheader {{
  color: #FFFFFF !important;
  font-weight: 900 !important;
  letter-spacing: 0.2px;
  font-size: 1.05rem !important;
  text-shadow: 0 1px 0 rgba(0,0,0,0.15);
}}

section[data-testid="stSidebar"] .stSubheader > div,
section[data-testid="stSidebar"] .stSubheader p {{
  color: #FFFFFF !important;
}}

section[data-testid="stSidebar"] hr {{
  border: none;
  border-top: 1px solid rgba(255,255,255,0.14);
  margin: 12px 0 14px 0;
}}

/* INPUT / TEXTAREA / NUMBER / DATE */
section[data-testid="stSidebar"] input,
section[data-testid="stSidebar"] textarea {{
  background: rgba(255,255,255,0.98) !important;
  color: #122230 !important;
  border: 1px solid rgba(255,255,255,0.22) !important;
  border-radius: 10px !important;
}}

section[data-testid="stSidebar"] input:focus,
section[data-testid="stSidebar"] textarea:focus {{
  border: 1px solid var(--azure) !important;
  box-shadow: 0 0 0 1px rgba(78,161,255,0.20) !important;
}}

/* SELECTBOX */
section[data-testid="stSidebar"] [data-baseweb="select"] > div {{
  background: rgba(255,255,255,0.98) !important;
  color: #122230 !important;
  border: 1px solid rgba(255,255,255,0.22) !important;
  border-radius: 10px !important;
}}

section[data-testid="stSidebar"] [data-baseweb="select"] span {{
  color: #122230 !important;
}}

/* MULTISELECT TAGS */
section[data-testid="stSidebar"] [data-baseweb="tag"] {{
  background: #DCEBFF !important;
  border-radius: 999px !important;
  color: #16324F !important;
  border: 1px solid #B8D5FF !important;
}}

/* CHECKBOX */
section[data-testid="stSidebar"] [data-testid="stCheckbox"] label {{
  color: #F8FBFF !important;
}}

/* BUTTONS SIDEBAR */
section[data-testid="stSidebar"] .stButton > button {{
  width: 100%;
  border-radius: 12px !important;
  border: 1px solid rgba(255,255,255,0.14) !important;
  background: linear-gradient(180deg, #21476E 0%, #183754 100%) !important;
  color: #FFFFFF !important;
  font-weight: 700 !important;
}}

section[data-testid="stSidebar"] .stButton > button:hover {{
  border-color: rgba(255,255,255,0.24) !important;
  background: linear-gradient(180deg, #27527F 0%, #1D4265 100%) !important;
  color: #FFFFFF !important;
}}

/* MAIN CARD */
.audit-card {{
  background: #ffffff;
  border: 1px solid rgba(16,24,40,0.08);
  border-radius: 18px;
  padding: 14px 16px;
  box-shadow: 0 6px 18px rgba(16,24,40,0.06);
}}

.badge {{
  display: inline-block;
  padding: 6px 10px;
  border-radius: 999px;
  font-weight: 800;
  font-size: 0.85rem;
  border: 1px solid rgba(16,24,40,0.12);
  background: var(--soft);
}}

.kpi {{
  font-size: 1.8rem;
  font-weight: 900;
  margin: 0;
}}

hr {{
  border: none;
  border-top: 1px solid var(--border);
  margin: 10px 0 14px 0;
}}

.status-box {{
  border-radius: 14px;
  padding: 10px 12px;
  margin: 4px 0 12px 0;
  border-left: 6px solid;
  font-weight: 700;
  font-size: 0.95rem;
}}

.status-pill {{
  display: inline-block;
  padding: 4px 10px;
  border-radius: 999px;
  font-weight: 800;
  font-size: 0.82rem;
  margin-bottom: 8px;
  border: 1px solid;
}}
</style>
""",
    unsafe_allow_html=True,
)

# -----------------------------
# PERCORSI DATI
# -----------------------------
BASE_DIR = Path(__file__).parent
LOGO_PATH = BASE_DIR / "assets" / "logo.png"
DATA_DIR = BASE_DIR / "data_certificazione"
CHECKLIST_PATH = BASE_DIR / "checklist_audit_iso.json"
TEMPLATE_DOCX_PATH = BASE_DIR / "Template_CheckList_MARKER.docx"
SAVED_AUDITS_DIR = BASE_DIR / "saved_audits"
AUDIT_HEADER_CSV = DATA_DIR / "audit_header.csv"
AUDIT_FINDINGS_CSV = DATA_DIR / "audit_findings.csv"
AUDIT_ACTIONS_CSV = DATA_DIR / "audit_actions.csv"


DEFAULT_CHECKLIST: Dict[str, Any] = {
    "meta": {
        "name": "Checklist ISO Integrata",
        "version": "0.1-demo",
        "description": "Base iniziale per audit ISO 9001 / 14001 / 45001.",
    },
    "sections": [
        {
            "code": "OPENING",
            "title": "Riunione iniziale",
            "type": "special",
            "items": [
                {
                    "id": "OPEN-01",
                    "clause": "Opening",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Sono stati presentati gruppo di audit, ruoli, scopo, modalità operative e canali di comunicazione?",
                },
                {
                    "id": "OPEN-02",
                    "clause": "Opening",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Sono stati confermati riservatezza, disponibilità guide, logistica, emergenze e partecipanti?",
                },
            ],
        },
        {
            "code": "4",
            "title": "4. Context of the organisation",
            "type": "requirements",
            "items": [
                {
                    "id": "4.1-01",
                    "clause": "4.1",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "L'organizzazione ha determinato questioni esterne e interne rilevanti per lo scopo aziendale e le direzioni strategiche?",
                },
                {
                    "id": "4.2-01",
                    "clause": "4.2",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "L'organizzazione ha determinato le parti interessate rilevanti e i relativi requisiti?",
                },
                {
                    "id": "4.3-01",
                    "clause": "4.3",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "L'ambito del sistema di gestione è determinato e documentato?",
                },
            ],
        },
        {
            "code": "5",
            "title": "5. Leadership",
            "type": "requirements",
            "items": [
                {
                    "id": "5.1-01",
                    "clause": "5.1",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "La direzione dimostra leadership e impegno verso il sistema di gestione?",
                },
                {
                    "id": "5.2-01",
                    "clause": "5.2",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Esiste una politica documentata pertinente, comunicata e compresa?",
                },
                {
                    "id": "5.3-01",
                    "clause": "5.3",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Ruoli, responsabilità e autorità sono determinati e compresi?",
                },
            ],
        },
        {
            "code": "6",
            "title": "6. Planning",
            "type": "requirements",
            "items": [
                {
                    "id": "6.1-01",
                    "clause": "6.1.1",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "L'organizzazione considera rischi e opportunità derivanti da contesto e parti interessate?",
                },
                {
                    "id": "6.1-02",
                    "clause": "6.1.2",
                    "norms": ["14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Sono determinati e documentati aspetti ambientali / pericoli SSL e relative azioni?",
                },
                {
                    "id": "6.2-01",
                    "clause": "6.2",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Obiettivi, responsabilità, risorse e tempi sono definiti?",
                },
            ],
        },
        {
            "code": "7",
            "title": "7. Support",
            "type": "requirements",
            "items": [
                {
                    "id": "7.1-01",
                    "clause": "7.1",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Risorse, persone, infrastrutture e ambiente operativo sono adeguatamente determinati e forniti?",
                },
                {
                    "id": "7.2-01",
                    "clause": "7.2",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Competenze, formazione e registrazioni sono disponibili e coerenti?",
                },
                {
                    "id": "7.5-01",
                    "clause": "7.5",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Le informazioni documentate sono disponibili, controllate e protette?",
                },
            ],
        },
        {
            "code": "8",
            "title": "8. Operation",
            "type": "requirements",
            "items": [
                {
                    "id": "8.1-01",
                    "clause": "8.1",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Pianificazione e controllo operativo risultano attuati e coerenti con i requisiti del SG?",
                },
                {
                    "id": "8.4-01",
                    "clause": "8.4",
                    "norms": ["9001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Processi, prodotti e servizi forniti dall'esterno sono controllati adeguatamente?",
                },
                {
                    "id": "8.1-EH-01",
                    "clause": "E&H 8.1/2",
                    "norms": ["14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Sono stabiliti, attuati e mantenuti processi per prepararsi e rispondere a potenziali emergenze?",
                },
            ],
        },
        {
            "code": "9",
            "title": "9. Performance evaluation",
            "type": "requirements",
            "items": [
                {
                    "id": "9.1-01",
                    "clause": "9.1.1",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "L'organizzazione valuta prestazioni ed efficacia del sistema di gestione?",
                },
                {
                    "id": "9.2-01",
                    "clause": "9.2",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Audit interni e relativo programma risultano pianificati, attuati e registrati?",
                },
                {
                    "id": "9.3-01",
                    "clause": "9.3",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Il riesame di direzione è svolto e documentato con input e output coerenti?",
                },
            ],
        },
        {
            "code": "10",
            "title": "10. Improvement",
            "type": "requirements",
            "items": [
                {
                    "id": "10.2-01",
                    "clause": "10.2",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Non conformità, azioni correttive e verifiche di efficacia sono gestite in modo strutturato?",
                },
                {
                    "id": "10.3-01",
                    "clause": "10.3",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "L'organizzazione promuove il miglioramento continuo del sistema di gestione?",
                },
            ],
        },
        {
            "code": "CLOSING",
            "title": "Riunione finale",
            "type": "special",
            "items": [
                {
                    "id": "CLOSE-01",
                    "clause": "Closing",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "Sono stati condivisi esiti audit, NC/OFI, raccomandazione e tempi per le azioni correttive?",
                },
                {
                    "id": "CLOSE-02",
                    "clause": "Closing",
                    "norms": ["9001", "14001", "45001"],
                    "stage_visibility": ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"],
                    "requirement": "È stata formalizzata la chiusura audit con conferma delle prossime fasi?",
                },
            ],
        },
    ],
}

ESITI = ["", "Conforme", "O", "NC", "Cm"]
ESITO_LABELS = {"": "Da compilare", "Conforme": "Conforme", "O": "Osservazione", "NC": "Non conformità", "Cm": "Commento"}
ESITO_UI = {
    "": {
        "label": "Da compilare",
        "icon": "🔵",
        "bg": "#EAF3FF",
        "border": "#9CC4FF",
        "text": "#1D4ED8",
    },
    "Conforme": {
        "label": "Conforme",
        "icon": "🟢",
        "bg": "#ECFDF3",
        "border": "#A7F3D0",
        "text": "#047857",
    },
    "O": {
        "label": "Osservazione",
        "icon": "🟡",
        "bg": "#FFF9DB",
        "border": "#FDE68A",
        "text": "#B45309",
    },
    "NC": {
        "label": "Non conformità",
        "icon": "🔴",
        "bg": "#FEF2F2",
        "border": "#FECACA",
        "text": "#B91C1C",
    },
    "Cm": {
        "label": "Commento",
        "icon": "🟠",
        "bg": "#FFF4E6",
        "border": "#FCD9BD",
        "text": "#C2410C",
    },
}
NORM_OPTIONS = ["9001", "14001", "45001"]
AUDIT_TYPE_OPTIONS = ["Iniziale", "1° Sorveglianza", "2° Sorveglianza", "Rinnovo"]
TEMPLATE_STAGE_OPTIONS = ["S1", "S2"]


# -----------------------------
# AUTH
# -----------------------------
def get_configured_users() -> Dict[str, str]:
    try:
        users = st.secrets["auth"]["users"]
    except Exception:
        return {}
    return {str(username): str(password) for username, password in users.items()}


def logout() -> None:
    st.session_state["authenticated"] = False
    st.session_state.pop("username", None)
    st.rerun()


def require_login() -> None:
    if st.session_state.get("authenticated", False):
        return

    users = get_configured_users()

    if not users:
        st.title("🔐 Accesso riservato")
        st.error("Autenticazione non configurata. Inserisci utenti e password nei secrets di Streamlit Cloud.")
        st.code(
            """[auth.users]
admin = "cambiaquesta_password"
cliente1 = "altra_password"""
        )
        st.stop()

    st.title("🔐 Accesso riservato")
    st.caption("Inserisci le credenziali per accedere all'app.")

    left, center, right = st.columns([1, 1.2, 1])
    with center:
        st.markdown(
            """
            <div class="audit-card" style="margin-top:8px;">
                <div style="font-size:1.05rem;font-weight:800;margin-bottom:6px;">Login</div>
                <div style="color:#475467;font-size:0.95rem;margin-bottom:10px;">
                    Accesso protetto tramite credenziali configurate nei secrets.
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        with st.form("login_form", clear_on_submit=False):
            username = st.text_input("Utente")
            password = st.text_input("Password", type="password")
            submitted = st.form_submit_button("Accedi", use_container_width=True)

        if submitted:
            username = username.strip()
            stored_password = users.get(username)
            if stored_password and hmac.compare_digest(password, stored_password):
                st.session_state["authenticated"] = True
                st.session_state["username"] = username
                st.session_state.pop("login_error", None)
                st.rerun()
            else:
                st.session_state["login_error"] = "Credenziali non valide."

        if st.session_state.get("login_error"):
            st.error(st.session_state["login_error"])

    st.stop()


# -----------------------------
# UTILITÀ DATI
# -----------------------------
def load_checklist() -> Dict[str, Any]:
    if CHECKLIST_PATH.exists():
        with CHECKLIST_PATH.open("r", encoding="utf-8") as f:
            return json.load(f)
    return deepcopy(DEFAULT_CHECKLIST)


def ensure_csv(path: Path, headers: List[str]) -> None:
    if not path.exists():
        with path.open("w", newline="", encoding="utf-8") as f:
            csv.writer(f, delimiter=";").writerow(headers)

def safe_read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    try:
        return pd.read_csv(path, sep=";", engine="python", on_bad_lines="skip")
    except Exception:
        return pd.DataFrame()

def sanitize_filename(value: str) -> str:
    safe_value = "".join(ch if ch.isalnum() or ch in ("_", "-", ".") else "_" for ch in value)
    return safe_value or "audit_certificazione"


def apply_loaded_audit_data(data: Dict[str, Any], checklist: Dict[str, Any]) -> bool:
    if not isinstance(data, dict):
        return False

    st.session_state.audit_header_iso = data.get("header", build_default_header())

    results: Dict[str, Dict[str, Any]] = {}
    for section in checklist["sections"]:
        for item in section["items"]:
            row = build_empty_result(item)
            row["section_code"] = section["code"]
            row["section_title"] = section["title"]
            results[item["id"]] = row

    saved_results = data.get("results", {}) or {}
    for key, value in saved_results.items():
        if key in results:
            results[key].update(value)

    st.session_state.audit_results_iso = results
    st.session_state.audit_actions_iso = data.get(
        "actions",
        [{"azione": "", "owner": "", "due_date": "", "status": "Aperta"}]
    )
    st.session_state.audit_obs_iso = data.get(
        "observations",
        {"strengths": "", "risks": "", "general_notes": ""}
    )

    for item_id, row in results.items():
        st.session_state[f"esito_{item_id}"] = row.get("esito", "")
        st.session_state[f"evidence_{item_id}"] = row.get("evidence", "")
        st.session_state[f"note_{item_id}"] = row.get("note", "")
        st.session_state[f"action_required_{item_id}"] = row.get("action_required", False)

    st.session_state["obs_strengths"] = st.session_state.audit_obs_iso.get("strengths", "")
    st.session_state["obs_risks"] = st.session_state.audit_obs_iso.get("risks", "")
    st.session_state["obs_general_notes"] = st.session_state.audit_obs_iso.get("general_notes", "")

    for idx, action in enumerate(st.session_state.audit_actions_iso):
        st.session_state[f"azione_{idx}"] = action.get("azione", "")
        st.session_state[f"owner_{idx}"] = action.get("owner", "")
        st.session_state[f"due_{idx}"] = action.get("due_date", "")
        st.session_state[f"status_{idx}"] = action.get("status", "Aperta")

    return True


def load_uploaded_audit(uploaded_file: Any, checklist: Dict[str, Any]) -> tuple[bool, str]:
    if uploaded_file is None:
        return False, "Nessun file selezionato."

    try:
        raw_bytes = uploaded_file.getvalue()
        data = json.loads(raw_bytes.decode("utf-8"))
    except UnicodeDecodeError:
        return False, "Il file JSON deve essere codificato in UTF-8."
    except json.JSONDecodeError:
        return False, "Il file caricato non è un JSON valido."
    except Exception as exc:
        return False, f"Impossibile leggere la bozza: {exc}"

    if not apply_loaded_audit_data(data, checklist):
        return False, "Struttura JSON non riconosciuta."

    return True, uploaded_file.name or "bozza_audit.json"

def build_default_header() -> Dict[str, Any]:
    return {
        "audit_id": "",
        "organization": "",
        "scope": "",
        "scope_changes": False,
        "scope_changes_years": "",
        "norms": ["9001"],
        "triennium": "2026-2028",
        "audit_type": "Iniziale",
        "template_stage_column": "S1",
        "audit_type_label": "Stage 1 / Stage 2 / Sorveglianza / Rinnovo",
        "audit_date": date.today().isoformat(),
        "audit_days": 1.0,
        "lead_auditor": "",
        "auditor_2": "",
        "auditor_3": "",
        "site_name": "",
        "references": "",
        "notes_methodology": "",
    }


def build_empty_result(item: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "item_id": item["id"],
        "section_code": "",
        "section_title": "",
        "clause": item.get("clause", ""),
        "requirement": item.get("requirement", ""),
        "evidence": "",
        "esito": "",
        "stage_mark": "",
        "note": "",
        "action_required": False,
    }


def init_session(checklist: Dict[str, Any]) -> None:
    if "audit_header_iso" not in st.session_state:
        st.session_state.audit_header_iso = build_default_header()

    if "audit_results_iso" not in st.session_state:
        results: Dict[str, Dict[str, Any]] = {}
        for section in checklist["sections"]:
            for item in section["items"]:
                row = build_empty_result(item)
                row["section_code"] = section["code"]
                row["section_title"] = section["title"]
                results[item["id"]] = row
        st.session_state.audit_results_iso = results

    if "audit_actions_iso" not in st.session_state:
        st.session_state.audit_actions_iso = [
            {"azione": "", "owner": "", "due_date": "", "status": "Aperta"}
        ]

    if "audit_obs_iso" not in st.session_state:
        st.session_state.audit_obs_iso = {
            "strengths": "",
            "risks": "",
            "general_notes": "",
        }


def reset_audit(checklist: Dict[str, Any]) -> None:
    for key in ["audit_header_iso", "audit_results_iso", "audit_actions_iso", "audit_obs_iso"]:
        if key in st.session_state:
            del st.session_state[key]
    init_session(checklist)


def visible_items(checklist: Dict[str, Any], selected_norms: List[str], stage: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    selected_norms_set = set(selected_norms)
    for section in checklist["sections"]:
        visible_section_items = []
        for item in section["items"]:
            item_norms = set(item.get("norms", []))
            item_stages = set(item.get("stage_visibility", []))
            if item_norms.intersection(selected_norms_set) and stage in item_stages:
                visible_section_items.append(item)
        if visible_section_items:
            rows.append({"section": section, "items": visible_section_items})
    return rows


def compute_score(results: Dict[str, Dict[str, Any]], visible_ids: List[str]) -> Dict[str, Any]:
    filtered = [results[i] for i in visible_ids if i in results]
    total = len(filtered)
    completed = sum(1 for r in filtered if (r.get("esito") or "").strip() and (r.get("evidence") or "").strip())
    pending = total - completed
    pct = round((completed / total) * 100, 1) if total else 0.0
    nc_minor = sum(1 for r in filtered if r.get("esito") == "NC")
    ofi = sum(1 for r in filtered if r.get("esito") == "O")
    conformi = sum(1 for r in filtered if r.get("esito") == "Conforme")
    commenti = sum(1 for r in filtered if r.get("esito") == "Cm")
    return {
        "total": total,
        "completed": completed,
        "pending": pending,
        "pct": pct,
        "nc_minor": nc_minor,
        "nc_major": 0,
        "ofi": ofi,
        "conformi": conformi,
        "commenti": commenti,
    }


def generate_audit_id(header: Dict[str, Any]) -> str:
    base = header.get("organization", "ORG").strip().upper().replace(" ", "")[:8] or "ORG"
    return f"{datetime.now().strftime('%y%m%d')}-{base}-CI"


def collect_payload(checklist: Dict[str, Any]) -> Dict[str, Any]:
    header = deepcopy(st.session_state.audit_header_iso)
    if not header.get("audit_id"):
        header["audit_id"] = generate_audit_id(header)

    results = deepcopy(st.session_state.audit_results_iso)
    actions = deepcopy(st.session_state.audit_actions_iso)
    observations = deepcopy(st.session_state.audit_obs_iso)

    visible = visible_items(checklist, header["norms"], header["audit_type"])
    for item_id in results:
        results[item_id]["stage_mark"] = header.get("template_stage_column", "S1")
    visible_ids = [item["id"] for group in visible for item in group["items"]]
    summary = compute_score(results, visible_ids)

    return {
        "header": header,
        "results": results,
        "actions": actions,
        "observations": observations,
        "visible_ids": visible_ids,
        "summary": summary,
        "checklist_meta": checklist.get("meta", {}),
    }


def build_export_filename(prefix: str, payload: Dict[str, Any], extension: str) -> str:
    audit_id = payload["header"].get("audit_id") or "draft"
    return f"{prefix}_{sanitize_filename(audit_id)}.{extension}"


def build_marker_payload(payload: Dict[str, Any]) -> Dict[str, str]:
    marker_map: Dict[str, str] = {}
    h = payload["header"]
    stage_col = h.get("template_stage_column", "S1")

    def add_text(key: str, value: str) -> None:
        value = (value or "").strip()
        if not value:
            return
        if marker_map.get(key):
            if value not in marker_map[key]:
                marker_map[key] += "\n" + value
        else:
            marker_map[key] = value

    severity_rank = {"": 0, "Conforme": 1, "Cm": 2, "O": 3, "NC": 4}

    marker_map["ID_AUDIT"] = h.get("audit_id", "")
    marker_map["ORGANIZZAZIONE"] = h.get("organization", "")
    marker_map["SCOPO"] = h.get("scope", "")
    marker_map["MODIFICHE_SCOPE_NO"] = "☒" if not h.get("scope_changes", False) else "☐"
    marker_map["MODIFICHE_SCOPE_SI"] = "☒" if h.get("scope_changes", False) else "☐"
    marker_map["ANNI_MODIFICHE_SCOPE"] = h.get("scope_changes_years", "")
    marker_map["ISO_9001"] = "☒" if "9001" in h.get("norms", []) else "☐"
    marker_map["ISO_14001"] = "☒" if "14001" in h.get("norms", []) else "☐"
    marker_map["ISO_45001"] = "☒" if "45001" in h.get("norms", []) else "☐"
    marker_map["TRIENNIO_1"] = h.get("triennium", "")
    marker_map["TRIENNIO_2"] = ""
    marker_map["TRIENNIO_3"] = ""
    marker_map["GG_UOMO_1"] = str(h.get("audit_days", ""))
    marker_map["GG_UOMO_2"] = ""
    marker_map["GG_UOMO_3"] = ""
    marker_map["NOME_1"] = h.get("lead_auditor", "")
    marker_map["NOME_2"] = h.get("auditor_2", "")
    marker_map["NOME_3"] = h.get("auditor_3", "")
    marker_map["TIPO_INIZIALE"] = "☒" if h.get("audit_type") == "Iniziale" else "☐"
    marker_map["TIPO_1_SORV"] = "☒" if h.get("audit_type") == "1° Sorveglianza" else "☐"
    marker_map["TIPO_2_SORV"] = "☒" if h.get("audit_type") == "2° Sorveglianza" else "☐"

    item_to_code = {
        "4.1-01": "4_1", "4.1-CC": "4_1", "4.2-01": "4_2", "4.2-CC": "4_2", "4.3-01": "4_3", "4.4-01": "Q_4_4_1", "4.4-02": "Q_4_4_2",
        "5.1-01": "Q_5_1_1_2", "5.2-01": "Q_5_2_1_2", "5.2-02": "Q_5_2_1_2", "5.3-01": "5_3", "5.4-01": "H_5_4", "5.4-ANX": "SSL_5_4",
        "6.1.1-01": "6_1_1", "6.1.1-SSL": "6_1_1", "6.1.2-01": "6_1_2", "6.1.2-ENV": "6_1_2", "6.1.2-HS": "6_1_2", "6.1.3-01": "EH_6_1_3", "6.1.4-01": "EH_6_1_4", "6.2-01": "6_2_1_2", "6.3-01": "Q_6_3",
        "7.1.1-01": "Q_7_1_1", "7.1.2-01": "Q_7_1_2", "7.1.3-01": "Q_7_1_3_4", "7.1.5-01": "Q_7_1_5_1", "7.1.5-02": "Q_7_1_5_2", "7.1.6-01": "Q_7_1_6", "7.2-01": "7_2", "7.3-01": "7_3", "7.4-01": "7_4", "7.4.2-01": "EH_7_4_2", "7.4.3-01": "EH_7_4_3", "7.5-01": "7_5_1",
        "8.1-Q-01": "8_1", "8.2.1-01": "8_2_1", "8.2.2-01": "8_2_2", "8.2.3-01": "8_2_3", "8.2.4-01": "8_2_4", "8.3.4-01": "8_3_4", "8.4.1-01": "8_4_1_8_4_2", "8.4.3-01": "8_4_3", "8.5.1-01": "8_5_1", "8.5.2-01": "8_5_2", "8.5.3-01": "8_5_3", "8.5.4-01": "8_5_4", "8.5.5-01": "8_5_5", "8.5.6-01": "8_5_6", "8.6-01": "8_6", "8.7-01": "8_7", "8EH-01": "EH_8_1_2", "8EH-02": "EH_8_1_2", "8EH-03": "EH_8_1_2", "8EH-04": "EH_8_1_2", "8.1.2-SSL": "SSL_8_1_2",
        "9.1.1-01": "9_1_1", "9.1.1-02": "9_1_1", "9.1.1-03": "9_1_1", "9.1.1-04": "9_1_1", "9.2-01": "9_2_1", "9.2-02": "9_2_2", "9.3-01": "Q_9_3_1_2_3", "9.3-02": "Q_9_3_1_2_3",
        "10.1-01": "10_1", "10.2.1-01": "Q_10_2_1", "10.2.1-SSL": "Q_10_2_1", "10.2.2-01": "Q_10_2_2", "10.3-01": "10_3",
    }

    esito_by_code: Dict[str, str] = {}
    for item_id in payload["visible_ids"]:
        row = payload["results"][item_id]
        code = item_to_code.get(item_id)
        if not code:
            continue
        add_text(f"AUDIT_EVIDENCE_{code}", row.get("evidence", ""))
        current = esito_by_code.get(code, "")
        new = row.get("esito", "")
        if severity_rank.get(new, 0) >= severity_rank.get(current, 0):
            esito_by_code[code] = new
        evidence_text = row.get("evidence", "")

        if stage_col == "S1":
            marker_map[f"S1_{code}"] = evidence_text
            marker_map[f"S2_{code}"] = ""
        else:
            marker_map[f"S1_{code}"] = ""
            marker_map[f"S2_{code}"] = evidence_text

    for code, esito in esito_by_code.items():
        marker_map[f"NCOFI_{code}"] = esito

    # Placeholder-safe defaults for markers present in template but not yet generated in app
    for key in ["CLOSING_SSL_RESP", "CLOSING_SSL_RESP_NOTE", "CLOSING_RLS", "CLOSING_RLS_NOTE"] + [f"CLOSING_{i:02d}" for i in range(1,17)] + [f"RECERT_{i:02d}" for i in range(1,16)]:
        marker_map.setdefault(key, "")

    return marker_map


def replace_markers_in_docx(template_path: Path, marker_map: Dict[str, str]) -> bytes:
    doc = Document(template_path)

    def replace_text(text: str) -> str:
        updated = text
        for key, value in marker_map.items():
            updated = updated.replace(f"{{{{{key}}}}}", value or "")
        return updated

    for para in doc.paragraphs:
        full = ''.join(run.text for run in para.runs) if para.runs else para.text
        new_text = replace_text(full)
        if new_text != full:
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ''
            else:
                para.text = new_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace_text(cell.text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def build_pdf(payload: Dict[str, Any]) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=14 * mm,
        rightMargin=14 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
        title="Audit Certificazione ISO",
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=15, leading=18, textColor=colors.HexColor(NAVY))
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=11, leading=14, textColor=colors.HexColor(NAVY))
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=8.8, leading=11, alignment=TA_LEFT)
    cell = ParagraphStyle("cell", parent=body, fontSize=8.1, leading=9.8)

    story = []
    story.append(Paragraph("Audit Certificazione ISO – Demo applicativa", h1))
    story.append(Spacer(1, 5))

    header = payload["header"]
    summary = payload["summary"]
    observations = payload["observations"]

    header_rows = [
        ["ID Audit", header.get("audit_id", "")],
        ["Organizzazione", header.get("organization", "")],
        ["Scopo", header.get("scope", "")],
        ["Norme", ", ".join(header.get("norms", []))],
        ["Triennio", header.get("triennium", "")],
        ["Tipo audit", f"{header.get('audit_type', '')} / {header.get('template_stage_column', 'S1')}"] ,
        ["Data audit", header.get("audit_date", "")],
        ["GG uomo", str(header.get("audit_days", ""))],
        ["Lead auditor", header.get("lead_auditor", "")],
        ["Sito / Organizzazione", header.get("site_name", "")],
        ["Avanzamento", f"{summary['completed']}/{summary['total']} ({summary['pct']}%)"],
    ]
    tbl = Table(header_rows, colWidths=[42 * mm, 140 * mm])
    tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(LIGHT)),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 8))

    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for item_id in payload["visible_ids"]:
        row = payload["results"][item_id]
        grouped.setdefault(row["section_title"], []).append(row)

    for section_title, rows in grouped.items():
        story.append(Paragraph(section_title, h2))
        story.append(Spacer(1, 3))
        table_rows = [["Clause", "Requirement", "Audit Evidence", "Esito", "Note"]]
        for r in rows:
            table_rows.append([
                Paragraph(r.get("clause", ""), cell),
                Paragraph(r.get("requirement", ""), cell),
                Paragraph(r.get("evidence", "-"), cell),
                Paragraph(r.get("esito", ""), cell),
                Paragraph(r.get("note", "-"), cell),
            ])
        t = Table(table_rows, colWidths=[18 * mm, 62 * mm, 62 * mm, 22 * mm, 26 * mm], repeatRows=1)
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(SOFT)),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("FONTSIZE", (0, 0), (-1, -1), 7.8),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(t)
        story.append(Spacer(1, 6))

    story.append(Paragraph("Osservazioni finali", h2))
    story.append(Paragraph(f"<b>Punti di forza:</b> {observations.get('strengths', '-')}", body))
    story.append(Paragraph(f"<b>Rischi / attenzioni:</b> {observations.get('risks', '-')}", body))
    story.append(Paragraph(f"<b>Note generali:</b> {observations.get('general_notes', '-')}", body))
    story.append(Spacer(1, 6))

    if payload["actions"]:
        story.append(Paragraph("Azioni conseguenti", h2))
        action_rows = [["Azione", "Owner", "Scadenza", "Stato"]]
        for action in payload["actions"]:
            if any(str(v).strip() for v in action.values()):
                action_rows.append([
                    Paragraph(action.get("azione", ""), cell),
                    Paragraph(action.get("owner", ""), cell),
                    Paragraph(action.get("due_date", ""), cell),
                    Paragraph(action.get("status", ""), cell),
                ])
        if len(action_rows) > 1:
            t = Table(action_rows, colWidths=[90 * mm, 35 * mm, 28 * mm, 28 * mm], repeatRows=1)
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(SOFT)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.0),
            ]))
            story.append(t)

    doc.build(story)
    buffer.seek(0)
    return buffer.read()

def build_word_report(payload: Dict[str, Any]) -> bytes:
    doc = Document()

    # Margini pagina
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    # Stili base
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Audit Certificazione ISO – Report")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(22, 50, 79)  # NAVY-ish

    doc.add_paragraph("")

    header = payload["header"]
    summary = payload["summary"]
    observations = payload["observations"]

    # Tabella dati generali
    doc.add_paragraph("Dati generali audit").runs[0].bold = True
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    header_rows = [
        ("ID Audit", header.get("audit_id", "")),
        ("Organizzazione", header.get("organization", "")),
        ("Scopo", header.get("scope", "")),
        ("Norme", ", ".join(header.get("norms", []))),
        ("Triennio", header.get("triennium", "")),
        ("Tipo audit", f"{header.get('audit_type', '')} / {header.get('template_stage_column', 'S1')}"),
        ("Data audit", header.get("audit_date", "")),
        ("GG uomo", str(header.get("audit_days", ""))),
        ("Lead auditor", header.get("lead_auditor", "")),
        ("Auditor 2", header.get("auditor_2", "")),
        ("Auditor 3", header.get("auditor_3", "")),
        ("Sito / Organizzazione", header.get("site_name", "")),
        ("Avanzamento", f"{summary['completed']}/{summary['total']} ({summary['pct']}%)"),
    ]

    for label, value in header_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(label)
        row_cells[1].text = str(value or "")

    doc.add_paragraph("")

    # Raggruppamento requisiti per sezione
    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for item_id in payload["visible_ids"]:
        row = payload["results"][item_id]
        grouped.setdefault(row["section_title"], []).append(row)

    for section_title, rows in grouped.items():
        p = doc.add_paragraph()
        r = p.add_run(section_title)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(22, 50, 79)

        sec_table = doc.add_table(rows=1, cols=5)
        sec_table.style = "Table Grid"
        hdr = sec_table.rows[0].cells
        hdr[0].text = "Clause"
        hdr[1].text = "Requirement"
        hdr[2].text = "Audit Evidence"
        hdr[3].text = "Esito"
        hdr[4].text = "Note"

        for riga in rows:
            cells = sec_table.add_row().cells
            cells[0].text = str(riga.get("clause", "") or "")
            cells[1].text = str(riga.get("requirement", "") or "")
            cells[2].text = str(riga.get("evidence", "-") or "-")
            cells[3].text = str(riga.get("esito", "") or "")
            cells[4].text = str(riga.get("note", "-") or "-")

        doc.add_paragraph("")

    # Osservazioni finali
    p = doc.add_paragraph()
    r = p.add_run("Osservazioni finali")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(22, 50, 79)

    doc.add_paragraph(f"Punti di forza: {observations.get('strengths', '-')}")
    doc.add_paragraph(f"Rischi / attenzioni: {observations.get('risks', '-')}")
    doc.add_paragraph(f"Note generali: {observations.get('general_notes', '-')}")

    doc.add_paragraph("")

    # Azioni
    actions_present = [a for a in payload["actions"] if any(str(v).strip() for v in a.values())]
    if actions_present:
        p = doc.add_paragraph()
        r = p.add_run("Azioni conseguenti")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(22, 50, 79)

        act_table = doc.add_table(rows=1, cols=4)
        act_table.style = "Table Grid"
        hdr = act_table.rows[0].cells
        hdr[0].text = "Azione"
        hdr[1].text = "Owner"
        hdr[2].text = "Scadenza"
        hdr[3].text = "Stato"

        for action in actions_present:
            cells = act_table.add_row().cells
            cells[0].text = str(action.get("azione", "") or "")
            cells[1].text = str(action.get("owner", "") or "")
            cells[2].text = str(action.get("due_date", "") or "")
            cells[3].text = str(action.get("status", "") or "")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

# -----------------------------
# UI HELPERS
# -----------------------------
def sidebar_header_form() -> None:
    current_user = st.session_state.get("username", "")

    if LOGO_PATH.exists():
        st.sidebar.image(str(LOGO_PATH), use_container_width=True)
        st.sidebar.markdown("<div style='height:8px;'></div>", unsafe_allow_html=True)

    st.sidebar.markdown(
        f"""
        <div style="
            padding:10px 12px;
            border-radius:14px;
            background: rgba(255,255,255,0.08);
            border:1px solid rgba(255,255,255,0.10);
            margin-bottom:12px;
        ">
            <div style="font-size:0.80rem; opacity:0.85; color:#DCEBFF;">Audit Manager</div>
            <div style="font-size:1.05rem; font-weight:800; color:white;">Configurazione audit</div>
            <div style="margin-top:8px; font-size:0.85rem; color:#DCEBFF;">Utente connesso: <b>{current_user}</b></div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if st.sidebar.button("🔓 Logout", key="logout_button"):
        logout()
    h = st.session_state.audit_header_iso
    st.sidebar.markdown(
        """
        <div style="
            color:#D4A63A;
            font-size:1.05rem;
            font-weight:900;
            margin:10px 0 8px 0;
            padding-bottom:6px;
            border-bottom:1px solid rgba(255,255,255,0.18);
        ">
            Dati audit
        </div>
        """,
        unsafe_allow_html=True,
    )
    h["audit_id"] = st.sidebar.text_input("ID Audit", value=h.get("audit_id", ""))
    h["organization"] = st.sidebar.text_input("Organizzazione", value=h.get("organization", ""))
    h["site_name"] = st.sidebar.text_input("Sito / Unità auditata", value=h.get("site_name", ""))
    h["scope"] = st.sidebar.text_area("Scopo", value=h.get("scope", ""), height=80)
    h["scope_changes"] = st.sidebar.checkbox("Modifiche scopo", value=h.get("scope_changes", False))
    h["scope_changes_years"] = st.sidebar.text_input("Anno/i modifiche", value=h.get("scope_changes_years", ""))
    h["norms"] = st.sidebar.multiselect("Norme applicabili", options=NORM_OPTIONS, default=h.get("norms", ["9001"])) or ["9001"]
    h["triennium"] = st.sidebar.text_input("Triennio di riferimento", value=h.get("triennium", "2026-2028"))
    h["audit_type"] = st.sidebar.selectbox("Tipo audit", AUDIT_TYPE_OPTIONS, index=AUDIT_TYPE_OPTIONS.index(h.get("audit_type", "Iniziale")))
    h["template_stage_column"] = st.sidebar.selectbox("Colonna template da popolare", TEMPLATE_STAGE_OPTIONS, index=TEMPLATE_STAGE_OPTIONS.index(h.get("template_stage_column", "S1")))
    h["audit_date"] = str(st.sidebar.date_input("Data audit", value=pd.to_datetime(h.get("audit_date", date.today().isoformat())).date()))
    h["audit_days"] = st.sidebar.number_input("GG. uomo", min_value=0.5, step=0.5, value=float(h.get("audit_days", 1.0)))
    h["lead_auditor"] = st.sidebar.text_input("Lead Auditor", value=h.get("lead_auditor", ""))
    h["auditor_2"] = st.sidebar.text_input("Auditor 2", value=h.get("auditor_2", ""))
    h["auditor_3"] = st.sidebar.text_input("Auditor 3", value=h.get("auditor_3", ""))
    h["references"] = st.sidebar.text_area("Riferimenti / note", value=h.get("references", ""), height=80)
    st.sidebar.markdown("---")
    st.sidebar.markdown(
        """
        <div style="
            color:#D4A63A;
            font-size:1.05rem;
            font-weight:900;
            margin:10px 0 8px 0;
            padding-bottom:6px;
            border-bottom:1px solid rgba(255,255,255,0.18);
        ">
            Bozze audit
        </div>
        """,
        unsafe_allow_html=True,
    )

    uploaded_draft = st.sidebar.file_uploader(
        "Carica bozza JSON",
        type=["json"],
        key="uploaded_audit_draft",
        help="Trascina qui una bozza JSON scaricata in precedenza per riprendere il lavoro.",
    )
    if st.sidebar.button("📂 Importa bozza JSON"):
        st.session_state["load_uploaded_audit"] = True

    st.sidebar.caption("In cloud la bozza non viene salvata sul server: per riprendere il lavoro scarica sempre il JSON audit sul tuo PC.")

    st.sidebar.markdown("---")
    st.sidebar.markdown(
        """
        <div style="
            color:#D4A63A;
            font-size:1.05rem;
            font-weight:900;
            margin:10px 0 8px 0;
            padding-bottom:6px;
            border-bottom:1px solid rgba(255,255,255,0.18);
        ">
            Legenda dei responsi forniti
        </div>
        """,
        unsafe_allow_html=True,
    )

    for esito_key in ["", "Conforme", "O", "NC", "Cm"]:
        ui = get_esito_ui(esito_key)
        st.sidebar.markdown(
            f"""
            <div class="status-pill" style="
                background:{ui['bg']};
                border-color:{ui['border']};
                color:{ui['text']};
                margin-bottom:6px;
            ">
                {ui['icon']} {ui['label']}
            </div>
            """,
            unsafe_allow_html=True,
        )

def get_esito_ui(esito: str) -> Dict[str, str]:
    return ESITO_UI.get(esito or "", ESITO_UI[""])

def render_summary(payload: Dict[str, Any]) -> None:
    s = payload["summary"]
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Requisiti visibili", s["total"])
    c2.metric("Compilati", s["completed"])
    c3.metric("Da completare", s["pending"])
    c4.metric("Osservazioni", s["ofi"])
    c5.metric("NC", s["nc_minor"])
    st.markdown(
        f"<div class='audit-card'><p class='kpi'>{s['pct']}%</p><div>Avanzamento compilazione audit. Colonna template selezionata: <b>{payload['header'].get('template_stage_column','S1')}</b></div></div>",
        unsafe_allow_html=True,
    )

def render_section(section: Dict[str, Any], items: List[Dict[str, Any]]) -> None:
    stage_col = st.session_state.audit_header_iso.get("template_stage_column", "S1")
    st.markdown(f"### {section['title']}")

    for item in items:
        row = st.session_state.audit_results_iso[item["id"]]
        row["stage_mark"] = stage_col

        esito_key = f"esito_{item['id']}"
        evidence_key = f"evidence_{item['id']}"
        note_key = f"note_{item['id']}"
        action_key = f"action_required_{item['id']}"

        current_esito = st.session_state.get(esito_key, row.get("esito", ""))
        row["esito"] = current_esito
        row["evidence"] = st.session_state.get(evidence_key, row.get("evidence", ""))
        row["note"] = st.session_state.get(note_key, row.get("note", ""))
        row["action_required"] = st.session_state.get(action_key, row.get("action_required", False))

        ui = get_esito_ui(current_esito)

        expander_title = (
            f"{ui['icon']} "
            f"{item.get('clause', '')} — {item['requirement']} "
            f"· {ui['label']}"
        )

        with st.expander(expander_title, expanded=False):
            st.markdown(
                f"""
                <div class="status-box" style="
                    background:{ui['bg']};
                    border-left-color:{ui['border']};
                    color:{ui['text']};
                ">
                    Stato attuale del requisito: {ui['icon']} {ui['label']}
                </div>
                """,
                unsafe_allow_html=True,
            )

            col1, col2 = st.columns([1, 1])

            row["esito"] = col1.selectbox(
                "Esito template",
                ESITI,
                index=ESITI.index(current_esito if current_esito in ESITI else ""),
                key=esito_key,
                format_func=lambda x: ESITO_LABELS.get(x, x),
            )

            # aggiorna UI dopo la scelta
            selected_ui = get_esito_ui(row["esito"])

            col2.markdown(
                f"""
                <div class="status-pill" style="
                    background:{selected_ui['bg']};
                    border-color:{selected_ui['border']};
                    color:{selected_ui['text']};
                ">
                    {selected_ui['icon']} {selected_ui['label']}
                </div>
                """,
                unsafe_allow_html=True,
            )
            col2.info(f"Scrittura colonna template: {stage_col}")

            row["evidence"] = st.text_area(
                "Audit Evidence",
                value=row.get("evidence", ""),
                key=f"evidence_{item['id']}",
                height=110,
            )

            row["note"] = st.text_area(
                "Note / commenti",
                value=row.get("note", ""),
                key=f"note_{item['id']}",
                height=80,
            )

            row["action_required"] = st.checkbox(
                "Richiede azione conseguente",
                value=bool(row.get("action_required", False)),
                key=f"action_required_{item['id']}",
            )

            st.caption(
                f"Norme: {', '.join(item.get('norms', []))} · "
                f"Clausola: {item.get('clause', '')} · "
                f"Marker stage: {stage_col}"
            )


def render_actions() -> None:
    st.markdown("### Azioni conseguenti")
    actions = st.session_state.audit_actions_iso
    for idx, action in enumerate(actions):
        c1, c2, c3, c4 = st.columns([4, 2, 2, 2])
        action["azione"] = c1.text_input("Azione", value=action.get("azione", ""), key=f"azione_{idx}")
        action["owner"] = c2.text_input("Owner", value=action.get("owner", ""), key=f"owner_{idx}")
        action["due_date"] = c3.text_input("Scadenza", value=action.get("due_date", ""), key=f"due_{idx}")
        action["status"] = c4.selectbox("Stato", ["Aperta", "In corso", "Chiusa"], index=["Aperta", "In corso", "Chiusa"].index(action.get("status", "Aperta")), key=f"status_{idx}")
    col_add, col_spacer = st.columns([1, 5])
    if col_add.button("➕ Aggiungi azione"):
        actions.append({"azione": "", "owner": "", "due_date": "", "status": "Aperta"})
        st.rerun()


def render_observations() -> None:
    st.markdown("### Osservazioni finali")
    obs = st.session_state.audit_obs_iso
    obs["strengths"] = st.text_area(
        "Punti di forza",
        value=obs.get("strengths", ""),
        height=90,
        key="obs_strengths"
    )
    obs["risks"] = st.text_area(
        "Rischi / attenzioni",
        value=obs.get("risks", ""),
        height=90,
        key="obs_risks"
    )
    obs["general_notes"] = st.text_area(
        "Note generali",
        value=obs.get("general_notes", ""),
        height=110,
        key="obs_general_notes"
    )


# -----------------------------
# APP
# -----------------------------
def main() -> None:
    require_login()

    checklist = load_checklist()
    init_session(checklist)
    sidebar_header_form()

    if st.session_state.get("load_uploaded_audit"):
        st.session_state.pop("load_uploaded_audit")
        uploaded_file = st.session_state.get("uploaded_audit_draft")
        ok, message = load_uploaded_audit(uploaded_file, checklist)
        if ok:
            st.success(f"Bozza caricata: {message}")
            st.rerun()
        st.error(message)

    st.title("📋 Audit Certificazione ISO")
    st.caption("Versione pronta per distribuzione cloud: nessun salvataggio persistente lato server, ripresa lavoro tramite upload della bozza JSON ed export finale in PDF / Word / DOCX template.")

    payload = collect_payload(checklist)
    render_summary(payload)

    st.markdown("---")
    visible = visible_items(checklist, payload["header"]["norms"], payload["header"]["audit_type"])

    tab1, tab2, tab3 = st.tabs(["Checklist audit", "Azioni", "Bozze e istruzioni"])

    with tab1:
        for group in visible:
            render_section(group["section"], group["items"])
        render_observations()

    with tab2:
        render_actions()

    with tab3:
        st.markdown("### Come gestire le bozze in distribuzione")
        st.info("Questa versione non salva audit sul server. Per riprendere un lavoro: 1) scarica il JSON audit, 2) conservalo sul tuo PC o su OneDrive/Drive, 3) ricaricalo nella sidebar con drag & drop quando vuoi continuare.")
        st.markdown("### Cosa scarica l’utente")
        st.markdown("""
- **JSON audit**: bozza completa da riaprire nell’app.
- **PDF demo**: report leggibile e stampabile.
- **WORD report**: report modificabile.
- **Marker map**: JSON tecnico per i placeholder.
- **DOCX template**: template Word compilato automaticamente, se presente nel repo.
""")

    st.markdown("---")
    payload = collect_payload(checklist)
    json_export = json.dumps(payload, ensure_ascii=False, indent=2)
    pdf_bytes = build_pdf(payload)
    word_bytes = build_word_report(payload)
    marker_payload_dict = build_marker_payload(payload)
    marker_payload = json.dumps(marker_payload_dict, ensure_ascii=False, indent=2)

    c1, c2, c3, c4, c5 = st.columns([1.35, 1, 1, 1, 1])

    c1.download_button(
        "⬇️ Scarica JSON audit",
        data=json_export,
        file_name=build_export_filename("audit_certificazione", payload, "json"),
        mime="application/json",
        help="Questo è il file da conservare per riaprire la bozza in un secondo momento.",
    )

    c2.download_button(
        "⬇️ Scarica PDF demo",
        data=pdf_bytes,
        file_name=build_export_filename("audit_certificazione_demo", payload, "pdf"),
        mime="application/pdf",
    )

    c3.download_button(
        "⬇️ Scarica WORD report",
        data=word_bytes,
        file_name=build_export_filename("audit_certificazione_report", payload, "docx"),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    c4.download_button(
        "⬇️ Scarica marker map",
        data=marker_payload,
        file_name=build_export_filename("marker_map", payload, "json"),
        mime="application/json",
    )

    if TEMPLATE_DOCX_PATH.exists():
        template_docx_bytes = replace_markers_in_docx(TEMPLATE_DOCX_PATH, marker_payload_dict)
        c5.download_button(
            "⬇️ Scarica DOCX template",
            data=template_docx_bytes,
            file_name=build_export_filename("audit_template", payload, "docx"),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        c5.warning("Template_CheckList_MARKER.docx non trovato")

    st.info(
        "Per distribuire la app tieni nel repo i file statici necessari, soprattutto checklist_audit_iso.json e Template_CheckList_MARKER.docx. Il salvataggio delle bozze resta in mano all’utente tramite download del JSON audit."
    )

    if st.button("🔄 Nuovo audit / reset"):
        reset_audit(checklist)
        st.rerun()


if __name__ == "__main__":
    main()
